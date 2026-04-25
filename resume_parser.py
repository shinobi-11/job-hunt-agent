"""Resume parsing from PDF and DOCX files."""
from pathlib import Path

import PyPDF2
from docx import Document


class ResumeParser:
    """Parse resume content from various formats."""

    @staticmethod
    def parse_resume(file_path: str) -> str:
        """Parse resume and extract text."""
        path = Path(file_path)

        if not path.exists():
            raise FileNotFoundError(f"Resume file not found: {file_path}")

        if path.suffix.lower() == ".pdf":
            return ResumeParser._parse_pdf(file_path)
        elif path.suffix.lower() in [".docx", ".doc"]:
            return ResumeParser._parse_docx(file_path)
        else:
            raise ValueError(f"Unsupported file format: {path.suffix}")

    @staticmethod
    def _parse_pdf(file_path: str) -> str:
        """Extract text from PDF resume."""
        text = []

        try:
            with open(file_path, "rb") as pdf_file:
                pdf_reader = PyPDF2.PdfReader(pdf_file)

                for page_num in range(len(pdf_reader.pages)):
                    page = pdf_reader.pages[page_num]
                    text.append(page.extract_text())

            return "\n".join(text)
        except Exception as e:
            raise ValueError(f"Error parsing PDF: {e}")

    @staticmethod
    def _parse_docx(file_path: str) -> str:
        """Extract text from DOCX resume."""
        text = []

        try:
            doc = Document(file_path)

            for para in doc.paragraphs:
                if para.text.strip():
                    text.append(para.text)

            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        if cell.text.strip():
                            text.append(cell.text)

            return "\n".join(text)
        except Exception as e:
            raise ValueError(f"Error parsing DOCX: {e}")

    @staticmethod
    def extract_sections(resume_text: str) -> dict:
        """Extract key sections from resume text."""
        sections = {
            "contact": [],
            "summary": [],
            "experience": [],
            "skills": [],
            "education": [],
            "certifications": []
        }

        lines = resume_text.split("\n")
        current_section = None

        section_keywords = {
            "contact": ["contact", "phone", "email", "linkedin"],
            "summary": ["summary", "objective", "profile"],
            "experience": ["experience", "work", "employment"],
            "skills": ["skills", "technical"],
            "education": ["education", "degree", "university"],
            "certifications": ["certifications", "certified"]
        }

        for line in lines:
            line_lower = line.lower().strip()

            # Detect section headers
            for section, keywords in section_keywords.items():
                if any(keyword in line_lower for keyword in keywords):
                    current_section = section
                    break

            if current_section and line.strip():
                sections[current_section].append(line.strip())

        return sections
