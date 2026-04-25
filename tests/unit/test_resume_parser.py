"""Unit tests for resume parser."""
from pathlib import Path

import pytest
from docx import Document

from resume_parser import ResumeParser


class TestResumeParser:
    def test_missing_file_raises(self):
        with pytest.raises(FileNotFoundError):
            ResumeParser.parse_resume("/nonexistent/resume.pdf")

    def test_unsupported_format_raises(self, tmp_path):
        unsupported = tmp_path / "resume.txt"
        unsupported.write_text("Plain text resume")
        with pytest.raises(ValueError, match="Unsupported"):
            ResumeParser.parse_resume(str(unsupported))

    def test_parses_docx(self, tmp_path):
        docx_path = tmp_path / "resume.docx"
        doc = Document()
        doc.add_paragraph("John Doe")
        doc.add_paragraph("Software Engineer")
        doc.add_paragraph("Python, React, AWS")
        doc.save(str(docx_path))

        text = ResumeParser.parse_resume(str(docx_path))
        assert "John Doe" in text
        assert "Software Engineer" in text
        assert "Python" in text

    def test_parses_docx_with_tables(self, tmp_path):
        docx_path = tmp_path / "resume.docx"
        doc = Document()
        doc.add_paragraph("Header")
        table = doc.add_table(rows=1, cols=2)
        table.rows[0].cells[0].text = "Skills"
        table.rows[0].cells[1].text = "Python, SQL"
        doc.save(str(docx_path))

        text = ResumeParser.parse_resume(str(docx_path))
        assert "Skills" in text
        assert "Python" in text

    def test_parses_real_resume_if_present(self):
        real = Path("./data/resume.pdf")
        if not real.exists():
            pytest.skip("Real resume not present")
        text = ResumeParser.parse_resume(str(real))
        assert len(text) > 100

    def test_extract_sections_detects_headers(self):
        resume = """
John Doe
SUMMARY
Experienced engineer

EXPERIENCE
Worked at Acme Corp

SKILLS
Python, React

EDUCATION
MIT 2020
"""
        sections = ResumeParser.extract_sections(resume)
        assert len(sections["summary"]) > 0
        assert len(sections["experience"]) > 0
        assert len(sections["skills"]) > 0
        assert len(sections["education"]) > 0
